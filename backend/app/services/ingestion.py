from __future__ import annotations

import json
import re
import shutil
import uuid
from dataclasses import dataclass
from pathlib import Path

from fastapi import UploadFile
from pydantic import BaseModel, Field

from app.core.config import settings


BASE_DIR = Path(__file__).resolve().parents[2]
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
MANIFEST_DIR = DATA_DIR / "manifests"


class ChunkRecord(BaseModel):
    chunk_id: str
    index: int
    text: str
    start_char: int
    end_char: int
    metadata: dict[str, str | int | list[str]] = Field(default_factory=dict)


class IngestionResult(BaseModel):
    document_id: str
    filename: str
    domain: str
    chunk_count: int
    status: str = "ingested"
    message: str
    source_path: str
    manifest_path: str
    chunks: list[ChunkRecord] = Field(default_factory=list)


@dataclass(slots=True)
class ParsedDocument:
    text: str
    source_type: str


def _ensure_data_dirs() -> None:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    MANIFEST_DIR.mkdir(parents=True, exist_ok=True)


def _normalize_text(text: str) -> str:
    cleaned = re.sub(r"\r\n?", "\n", text)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    return cleaned.strip()


def _read_text_file(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def _extract_pdf_text(file_path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise RuntimeError("PDF ingestion requires PyMuPDF to be installed.") from exc

    doc = fitz.open(file_path)
    
    # Step 1: Sample font sizes to find standard body font size and list of heading sizes
    font_sizes = []
    sample_pages = min(10, len(doc))
    for page_idx in range(sample_pages):
        page = doc[page_idx]
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" in b:
                for l in b["lines"]:
                    for s in l["spans"]:
                        if s["text"].strip():
                            # Round to 0.5 to smooth out slight variations
                            font_sizes.append(round(s["size"] * 2) / 2)
                            
    if not font_sizes:
        # Fallback to standard text extraction
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text("text") or "")
        return "\n".join(pages)
        
    from collections import Counter
    size_counts = Counter(font_sizes)
    body_size = size_counts.most_common(1)[0][0]
    
    # Threshold: heading must be at least 12% larger than body text (e.g. 10.64pt for 9.5pt body)
    heading_threshold = body_size * 1.12
    larger_sizes = sorted([sz for sz in size_counts.keys() if sz >= heading_threshold], reverse=True)
    
    pages: list[str] = []
    total_pages = len(doc)
    
    for page_idx, page in enumerate(doc):
        page_height = page.rect.height
        raw_text = page.get_text("text") or ""
        page_text_lower = raw_text.lower()
        
        # Heuristic 1: Skip Table of Contents / Front Matter pages
        if page_idx < min(15, total_pages - 1):
            is_toc = (
                "table of contents" in page_text_lower or
                "contents" in page_text_lower.split("\n")[:5] or
                page_text_lower.count("...") > 2
            )
            if is_toc:
                continue
                
        # Heuristic 2: Skip Index / Back Matter pages at the end of the book
        if page_idx >= max(0, total_pages - 15):
            is_index = (
                "index" in page_text_lower.split("\n")[:3] or
                ("index" in page_text_lower and re.search(r"\w+,\s*\d+", raw_text) is not None)
            )
            if is_index:
                continue

        blocks = page.get_text("dict")["blocks"]
        page_lines: list[str] = []
        
        for b in blocks:
            if "lines" in b:
                block_lines: list[str] = []
                weighted_size_sum = 0
                total_chars = 0
                is_bold = False
                is_mono = False
                
                for l in b["lines"]:
                    line_text = ""
                    for s in l["spans"]:
                        span_text = s["text"]
                        char_count = len(span_text)
                        if char_count > 0:
                            line_text += span_text
                            weighted_size_sum += s["size"] * char_count
                            total_chars += char_count
                            
                        font_lower = s["font"].lower()
                        if s["flags"] & 2 or "bold" in font_lower:
                            is_bold = True
                        if any(term in font_lower for term in ["mono", "cour", "consolas", "fira", "code"]):
                            is_mono = True
                            
                    block_lines.append(line_text)
                    
                block_content = "\n".join(block_lines).strip()
                if not block_content:
                    continue
                
                # Calculate character-length weighted average size of the block
                avg_size = weighted_size_sum / total_chars if total_chars > 0 else 0
                rounded_size = round(avg_size * 2) / 2
                
                # Check if it should be treated as a markdown heading
                # Monospace (code) blocks are never headings
                is_heading = (
                    len(block_content) < 120 and 
                    rounded_size >= heading_threshold and 
                    rounded_size in larger_sizes and
                    not is_mono
                )
                
                # Heuristic 3: Strip Header / Footer blocks based on bounding box y-coordinates
                # We exempt headings from being treated as margin blocks
                x0, y0, x1, y1 = b["bbox"]
                is_margin_block = (y0 < 50 or y1 > page_height - 85) and (not is_heading)
                
                # If block is in the margin and looks like a header/footer/copyright/page number, skip it
                if is_margin_block and (
                    len(block_content) < 100 or
                    block_content.isdigit() or
                    "page" in block_content.lower() or
                    "©" in block_content or
                    "copyright" in block_content.lower()
                ):
                    continue
                
                if is_heading:
                    rank = larger_sizes.index(rounded_size)
                    if rank == 0:
                        block_content = f"# {block_content}"
                    elif rank == 1:
                        block_content = f"## {block_content}"
                    else:
                        block_content = f"### {block_content}"
                elif len(block_content) < 120 and rounded_size == body_size and is_bold and not is_mono:
                    # Bold text at body size is H3 (e.g. inline subsection labels)
                    block_content = f"### {block_content}"
                elif is_mono:
                    # Wrap monospace code blocks in markdown fences
                    block_content = f"```python\n{block_content}\n```"
                            
                page_lines.append(block_content)
                
        pages.append("\n\n".join(page_lines))
        
    combined_text = "\n\n".join(pages)
    
    # Heuristic 4: Merge hyphenated word splits across line breaks
    combined_text = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", combined_text)
    
    return combined_text


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency issue
        raise RuntimeError("DOCX ingestion requires python-docx to be installed.") from exc

    document = Document(str(file_path))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name and style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
                text = f"{'#' * level} {text}"
            except (ValueError, IndexError):
                text = f"# {text}"
                
        paragraphs.append(text)
        
    return "\n".join(paragraphs)


def parse_document(file_path: Path) -> ParsedDocument:
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".rst"}:
        text = _read_text_file(file_path.read_bytes())
        return ParsedDocument(text=_normalize_text(text), source_type="text")
    if suffix == ".pdf":
        return ParsedDocument(text=_normalize_text(_extract_pdf_text(file_path)), source_type="pdf")
    if suffix == ".docx":
        return ParsedDocument(text=_normalize_text(_extract_docx_text(file_path)), source_type="docx")
    raise ValueError(f"Unsupported file type: {suffix}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> list[ChunkRecord]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    # 1. Identify all headers for metadata tracking
    header_pattern = re.compile(r"^(#+)\s+(.+)$", re.MULTILINE)
    headers = []
    for match in header_pattern.finditer(normalized):
        headers.append({
            "level": len(match.group(1)),
            "title": match.group(2).strip(),
            "start": match.start(),
            "end": match.end()
        })

    chunks: list[ChunkRecord] = []
    start = 0
    index = 0
    total_length = len(normalized)

    # 2. Split with priority on headers and whitespace
    # We use a greedy approach but try to find the best split point
    separators = [
        "\n# ", "\n## ", "\n### ", "\n#### ",
        "\n\n", "\n", " ", ""
    ]

    while start < total_length:
        if start + chunk_size >= total_length:
            end = total_length
        else:
            end = start + chunk_size
            found_split = False
            
            # Look for the best separator to split on, searching backwards from 'end'
            for sep in separators:
                # We want to split at a separator that is not too far back
                # Limit search to the last 50% of the chunk size to avoid tiny chunks
                search_start = max(start + overlap, start + (chunk_size // 2))
                split_point = normalized.rfind(sep, search_start, end)
                
                if split_point != -1:
                    if sep.strip().startswith("#"):
                        # For headers, we prefer to split right BEFORE the newline that precedes the header
                        # or just before the header itself if it's the start of the string.
                        end = split_point
                    else:
                        # For other separators, split after the separator
                        end = split_point + len(sep)
                    found_split = True
                    break
            
            # If no good separator found, we just take the chunk_size as is

        chunk_content = normalized[start:end].strip()

        # 3. Determine the header context (breadcrumb) for this chunk
        # Find the active header path at the 'start' of this chunk
        breadcrumb = []
        path = {}
        for h in headers:
            if h["start"] <= start:
                path[h["level"]] = h["title"]
                # Clear any sub-levels that are no longer active
                for level in list(path.keys()):
                    if level > h["level"]:
                        del path[level]
            else:
                break
        
        sorted_levels = sorted(path.keys())
        breadcrumb = [path[l] for l in sorted_levels]

        if chunk_content:
            chunks.append(
                ChunkRecord(
                    chunk_id=str(uuid.uuid4()),
                    index=index,
                    text=chunk_content,
                    start_char=start,
                    end_char=end,
                    metadata={
                        "headers": breadcrumb,
                        "header_context": " > ".join(breadcrumb) if breadcrumb else ""
                    }
                )
            )
            index += 1

        if end >= total_length:
            break

        # Move start forward, ensuring we make progress and respect overlap
        # If we split at a header, we might want less overlap to avoid repeating the header too much
        start = end - overlap
        if start < 0:
            start = 0
        
        # Safety: if we didn't move forward, force progress
        if start <= chunks[-1].start_char:
            start = end

    return chunks


async def ingest_document(file: UploadFile, domain: str) -> IngestionResult:
    _ensure_data_dirs()

    document_id = str(uuid.uuid4())
    original_filename = file.filename or "uploaded-document"
    safe_filename = Path(original_filename).name
    document_dir = UPLOAD_DIR / document_id
    document_dir.mkdir(parents=True, exist_ok=True)

    source_path = document_dir / safe_filename
    with source_path.open("wb") as target_file:
        shutil.copyfileobj(file.file, target_file)

    parsed_document = parse_document(source_path)
    chunks = chunk_text(parsed_document.text)

    # Index chunks in the local ChromaDB vector store
    from app.services.vector_store import index_document_chunks
    await index_document_chunks(
        document_id=document_id,
        filename=safe_filename,
        domain=domain or settings.default_domain,
        chunks=chunks
    )

    manifest = {
        "document_id": document_id,
        "filename": safe_filename,
        "domain": domain or settings.default_domain,
        "source_type": parsed_document.source_type,
        "source_path": str(source_path),
        "chunk_count": len(chunks),
        "chunks": [chunk.model_dump() for chunk in chunks],
    }

    manifest_path = MANIFEST_DIR / f"{document_id}.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    return IngestionResult(
        document_id=document_id,
        filename=safe_filename,
        domain=domain or settings.default_domain,
        chunk_count=len(chunks),
        message="Document ingested successfully.",
        source_path=str(source_path),
        manifest_path=str(manifest_path),
        chunks=chunks,
    )


# --- Document Management Functions ---

async def get_all_documents() -> list[dict]:
    """List all uploaded documents from the manifest directory."""
    if not MANIFEST_DIR.exists():
        return []
    
    documents = []
    for manifest_path in MANIFEST_DIR.glob("*.json"):
        try:
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            documents.append({
                "document_id": data.get("document_id"),
                "filename": data.get("filename"),
                "domain": data.get("domain"),
                "source_type": data.get("source_type"),
                "chunk_count": data.get("chunk_count"),
                "source_path": data.get("source_path"),
                "manifest_path": str(manifest_path),
            })
        except Exception:
            continue
    return documents


async def get_document(document_id: str) -> dict | None:
    """Get a specific document by ID."""
    manifest_path = MANIFEST_DIR / f"{document_id}.json"
    if not manifest_path.exists():
        return None
    try:
        data = json.loads(manifest_path.read_text(encoding="utf-8"))
        return data
    except Exception:
        return None


async def delete_document(document_id: str) -> bool:
    """Delete a document and all its chunks from the system."""
    from app.services.vector_store import delete_document_chunks
    
    manifest_path = MANIFEST_DIR / f"{document_id}.json"
    if not manifest_path.exists():
        return False
    
    try:
        data = json.loads(manifest_path.read_text(encoding="utf-8"))
        domain = data.get("domain", settings.default_domain)
        
        # Extract chunk IDs from manifest for reliable deletion
        chunk_ids = []
        for chunk in data.get("chunks", []):
            if isinstance(chunk, dict) and "chunk_id" in chunk:
                chunk_ids.append(chunk["chunk_id"])
        
        # Delete from ChromaDB vector store (by chunk IDs if available)
        await delete_document_chunks(document_id, domain, chunk_ids=chunk_ids)
        
        # Delete manifest
        manifest_path.unlink(missing_ok=True)
        
        # Delete uploaded source files
        document_dir = UPLOAD_DIR / document_id
        if document_dir.exists():
            shutil.rmtree(document_dir)
        
        return True
    except Exception as exc:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Failed to delete document {document_id}: {exc}")
        return False


async def delete_all_documents() -> dict:
    """Delete ALL documents, chunks, manifests, and source files."""
    from app.services.vector_store import get_chroma_client, delete_document_chunks
    
    deleted = 0
    failed = 0
    
    for manifest_path in MANIFEST_DIR.glob("*.json"):
        try:
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            doc_id = data.get("document_id")
            domain = data.get("domain", settings.default_domain)
            
            # Extract chunk IDs from manifest for reliable deletion
            chunk_ids = []
            for chunk in data.get("chunks", []):
                if isinstance(chunk, dict) and "chunk_id" in chunk:
                    chunk_ids.append(chunk["chunk_id"])
            
            # Delete from ChromaDB
            await delete_document_chunks(doc_id, domain, chunk_ids=chunk_ids)
            
            # Delete manifest
            manifest_path.unlink(missing_ok=True)
            
            # Delete source files
            doc_dir = UPLOAD_DIR / doc_id
            if doc_dir.exists():
                shutil.rmtree(doc_dir)
            
            deleted += 1
        except Exception as exc:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to delete document during cleanup: {exc}")
            failed += 1
            manifest_path.unlink(missing_ok=True)
    
    # Clean up empty ChromaDB collections
    try:
        client = get_chroma_client()
        for collection in client.list_collections():
            try:
                coll = client.get_collection(collection.name)
                if coll.count() == 0:
                    client.delete_collection(collection.name)
            except Exception:
                pass
    except Exception:
        pass
    
    return {"deleted": deleted, "failed": failed, "message": f"Cleaned up {deleted} documents."}
